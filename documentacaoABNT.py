# ==========================================
# GERAÇÃO DE DOCUMENTOS WORD (ABNT)
# Equivalente em Python das funções do server.js original
# ==========================================
import io
import unicodedata

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONTE = "Arial"
TAMANHO_TEXTO = Pt(12)

SECOES_RESUMO = {"resumo", "abstract"}
SECOES_REFERENCIAS = {
    "referencias",
    "referencia bibliografica",
    "referencias bibliograficas",
}

INDENT_PRIMEIRA_LINHA = Cm(1.25)

MARGENS = {
    "top": Cm(3),
    "left": Cm(3),
    "bottom": Cm(2),
    "right": Cm(2),
}


def normalizar(txt):
    """Remove acentos, deixa minúsculo e sem espaços nas pontas."""
    txt = unicodedata.normalize("NFKD", txt)
    txt = "".join(c for c in txt if not unicodedata.combining(c))
    return txt.lower().strip()


# ---------- helpers de baixo nível ----------

def _aplicar_fonte(run, tamanho=TAMANHO_TEXTO, negrito=False, italico=False):
    run.font.name = FONTE
    run.font.size = tamanho
    run.font.bold = negrito
    run.font.italic = italico
    # garante a fonte também para caracteres de script complexo (evita fallback)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONTE)


def _paragrafo_pagebreak(doc):
    """Paragrafo vazio que força quebra de página (equivalente a pageBreakBefore)."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


def p_centralizado(doc, texto, tamanho=TAMANHO_TEXTO, negrito=False, espaco_antes=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(espaco_antes)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(texto)
    _aplicar_fonte(run, tamanho=tamanho, negrito=negrito)
    return p


def p_bloco_direita(doc, texto, largura_cm=7.5, tamanho=TAMANHO_TEXTO):
    recuo_esquerdo_cm = 15 - largura_cm - 2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(recuo_esquerdo_cm)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(texto)
    _aplicar_fonte(run, tamanho=tamanho)
    return p


def paragrafos_vazios(doc, n):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0


# ---------- cabeçalho com número de página ----------

def _adicionar_campo_pagina(paragraph):
    run = paragraph.add_run()
    _aplicar_fonte(run)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def configurar_pagina_e_cabecalho(doc, com_cabecalho_pagina):
    section = doc.sections[0]
    section.top_margin = MARGENS["top"]
    section.left_margin = MARGENS["left"]
    section.bottom_margin = MARGENS["bottom"]
    section.right_margin = MARGENS["right"]

    if com_cabecalho_pagina:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _adicionar_campo_pagina(p)


# ---------- natureza do trabalho (folha de rosto) ----------

def natureza_texto(tipo_documento, curso, instituicao, orientador):
    txt = ""
    if tipo_documento == "TCC":
        txt = (
            f"Trabalho de Conclusão de Curso apresentado ao curso de {curso or '[curso]'} "
            f"da {instituicao}, como requisito parcial para obtenção do título de graduado."
        )
    elif tipo_documento == "Trabalho Acadêmico":
        txt = (
            f"Trabalho apresentado à disciplina de {curso or '[disciplina]'}, "
            f"do curso de graduação da {instituicao}, como requisito parcial de avaliação."
        )
    elif tipo_documento == "Relatório de Estágio":
        txt = (
            f"Relatório de Estágio Supervisionado apresentado ao curso de {curso or '[curso]'} "
            f"da {instituicao}, como requisito parcial para aprovação na disciplina "
            f"de Estágio Supervisionado."
        )
    if orientador and orientador.strip():
        txt += f"\n\nOrientador(a): {orientador.strip()}"
    return txt


# ---------- elementos pré-textuais ----------

def adicionar_capa(doc, instituicao, nome_aluno, titulo_trabalho, cidade, ano):
    p_centralizado(doc, instituicao.upper())
    paragrafos_vazios(doc, 6)
    p_centralizado(doc, nome_aluno.upper())
    paragrafos_vazios(doc, 6)
    p_centralizado(doc, titulo_trabalho.upper(), negrito=True)
    paragrafos_vazios(doc, 10)
    p_centralizado(doc, cidade)
    p_centralizado(doc, str(ano))
    _paragrafo_pagebreak(doc)


def adicionar_folha_rosto(doc, nome_aluno, titulo_trabalho, instituicao, tipo_documento, curso, orientador, cidade, ano):
    p_centralizado(doc, nome_aluno.upper())
    paragrafos_vazios(doc, 6)
    p_centralizado(doc, titulo_trabalho.upper(), negrito=True)
    paragrafos_vazios(doc, 6)

    natureza = natureza_texto(tipo_documento, curso, instituicao, orientador)
    for bloco in natureza.split("\n\n"):
        p_bloco_direita(doc, bloco)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0

    paragrafos_vazios(doc, 8)
    p_centralizado(doc, cidade)
    p_centralizado(doc, str(ano))
    _paragrafo_pagebreak(doc)


def adicionar_cabecalho_artigo(doc, titulo_trabalho, nome_aluno, instituicao):
    p_centralizado(doc, titulo_trabalho.upper(), negrito=True)
    paragrafos_vazios(doc, 1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(nome_aluno)
    _aplicar_fonte(run)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.paragraph_format.line_spacing = 1.0
    run2 = p2.add_run(instituicao)
    _aplicar_fonte(run2, tamanho=Pt(10))

    paragrafos_vazios(doc, 1)


def montar_elementos_pretextuais(doc, tipo_documento, nome_aluno, instituicao, titulo_trabalho, curso, orientador, cidade, ano):
    if tipo_documento in ("TCC", "Trabalho Acadêmico", "Relatório de Estágio"):
        adicionar_capa(doc, instituicao, nome_aluno, titulo_trabalho, cidade, ano)
        adicionar_folha_rosto(doc, nome_aluno, titulo_trabalho, instituicao, tipo_documento, curso, orientador, cidade, ano)
    elif tipo_documento == "Artigo Científico":
        adicionar_cabecalho_artigo(doc, titulo_trabalho, nome_aluno, instituicao)


# ---------- corpo do documento (markdown simplificado -> parágrafos ABNT) ----------

def adicionar_corpo(doc, texto_ia):
    linhas = texto_ia.strip().split("\n")
    modo = "normal"

    for linha_bruta in linhas:
        linha = linha_bruta.strip()
        if not linha:
            continue

        if linha.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(linha[4:].strip())
            _aplicar_fonte(run, negrito=True, italico=True)
            continue

        if linha.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(linha[3:].strip())
            _aplicar_fonte(run, negrito=True)
            continue

        if linha.startswith("# "):
            titulo = linha[2:].strip()
            norm = normalizar(titulo)
            if norm in SECOES_RESUMO:
                modo = "resumo"
            elif norm in SECOES_REFERENCIAS:
                modo = "referencias"
            else:
                modo = "normal"
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(titulo.upper())
            _aplicar_fonte(run, negrito=True)
            continue

        if modo == "resumo":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(linha)
            _aplicar_fonte(run)
        elif modo == "referencias":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(linha)
            _aplicar_fonte(run)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = INDENT_PRIMEIRA_LINHA
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(linha)
            _aplicar_fonte(run)


# ---------- montagem final ----------

def gerar_documento_word_bytes(texto_ia, tipo_documento, nome_aluno, instituicao, titulo_trabalho,
                                curso="", orientador="", cidade="Cidade", ano="2026"):
    doc = Document()
    configurar_pagina_e_cabecalho(doc, com_cabecalho_pagina=True)
    montar_elementos_pretextuais(doc, tipo_documento, nome_aluno, instituicao, titulo_trabalho,
                                  curso, orientador, cidade, ano)
    adicionar_corpo(doc, texto_ia)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def gerar_documento_simples_bytes(texto):
    doc = Document()
    configurar_pagina_e_cabecalho(doc, com_cabecalho_pagina=False)
    adicionar_corpo(doc, texto)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()